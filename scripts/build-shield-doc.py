#!/usr/bin/env python3
"""
Shield/Envelop branded Word document builder.
Rex calls this via: python3 scripts/build-shield-doc.py --title "..." --content "..." --out "/tmp/out.docx"
Content is plain text — sections separated by blank lines, section headers in ALL CAPS or ending with ':'.
Tables are passed as JSON lines: TABLE_JSON:[{"headers":["A","B"],"rows":[["1","2"]]}]
"""

import sys, os, json, argparse, tempfile
from docx import Document
from docx.shared import Pt, RGBColor, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── Brand colours ──────────────────────────────────────────────────────────────
BLUE  = RGBColor(0x0E, 0xA5, 0xE9)
GRAY  = RGBColor(0x6B, 0x72, 0x80)
DARK  = RGBColor(0x1F, 0x29, 0x37)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FOOTER_TEXT = (
    "Shield Technologies Corporation  |  envelopcovers.com  |  "
    "andy.parks@shieldtechnologies.com  |  740.219.1122"
)

LOGO_PATH = os.path.join(os.path.dirname(__file__), "..", "..", ".openclaw", "workspace", "documents", "envelop-logo-extracted.png")
# Fallback to workspace path
if not os.path.exists(LOGO_PATH):
    LOGO_PATH = "/root/.openclaw/workspace/documents/envelop-logo-extracted.png"


def build_doc(title: str, subtitle: str, content: str, out_path: str):
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin  = Emu(914400)
    sec.right_margin = Emu(914400)
    sec.top_margin   = Emu(914400)
    sec.bottom_margin= Emu(731520)

    # ── Footer ─────────────────────────────────────────────────────────────────
    footer_p = sec.footer.paragraphs[0]
    footer_p.text = FOOTER_TEXT
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in footer_p.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY

    # ── Header: logo left, wordmark right ──────────────────────────────────────
    header = sec.header
    ht = header.add_table(rows=1, cols=2, width=Emu(10800000))
    ht.autofit = True
    lc, rc = ht.rows[0].cells
    for cell in (lc, rc):
        tcPr = cell._tc.get_or_add_tcPr()
        borders = tcPr.makeelement(qn("w:tcBorders"), {})
        for edge in ("top","left","bottom","right"):
            e = borders.makeelement(qn(f"w:{edge}"), {qn("w:val"): "none"})
            borders.append(e)
        tcPr.append(borders)

    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_PATH):
        lp.add_run().add_picture(LOGO_PATH, width=Inches(1.8))
    else:
        r = lp.add_run("ENVELOP®")
        r.bold = True; r.font.size = Pt(14); r.font.color.rgb = BLUE

    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run("ENVELOP PROTECTIVE COVERS")
    rr.bold = True; rr.font.size = Pt(13); rr.font.color.rgb = BLUE

    # ── Helpers ────────────────────────────────────────────────────────────────
    def add_title(text, size=16, color=BLUE, bold=True, space_after=4):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color
        p.paragraph_format.space_after = Pt(space_after)

    def add_section(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(12); r.font.color.rgb = BLUE
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(6)

    def add_body(text, bold_label=None):
        p = doc.add_paragraph()
        if bold_label:
            r = p.add_run(bold_label)
            r.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK
        r2 = p.add_run(text)
        r2.font.size = Pt(10); r2.font.color.rgb = DARK

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            r = p.add_run(bold_prefix)
            r.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK
        r2 = p.add_run(text)
        r2.font.size = Pt(10); r2.font.color.rgb = DARK

    def add_divider():
        p = doc.add_paragraph("─" * 80)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(7); r.font.color.rgb = GRAY

    def add_table(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = ""
            p = hdr[i].paragraphs[0]
            r = p.add_run(h)
            r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE
            shd = hdr[i]._tc.get_or_add_tcPr()
            shd.append(shd.makeelement(qn("w:shd"), {qn("w:fill"): "0EA5E9"}))
        for ridx, row in enumerate(rows):
            cells = t.add_row().cells
            is_last = ridx == len(rows) - 1
            for i, val in enumerate(row):
                cells[i].text = ""
                p = cells[i].paragraphs[0]
                r = p.add_run(str(val))
                r.font.size = Pt(9); r.bold = is_last; r.font.color.rgb = DARK
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Title block ────────────────────────────────────────────────────────────
    add_title(title.upper(), size=16)
    if subtitle:
        add_title(subtitle, size=13, color=DARK, bold=True, space_after=2)

    # ── Parse and render content ────────────────────────────────────────────────
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # TABLE_JSON marker
        if stripped.startswith("TABLE_JSON:"):
            try:
                table_data = json.loads(stripped[11:])
                add_table(table_data["headers"], table_data["rows"])
            except Exception as e:
                add_body(f"[Table parse error: {e}]")
            i += 1
            continue

        # Divider
        if stripped in ("---", "───", ""):
            if stripped == "---":
                add_divider()
            i += 1
            continue

        # Section header: ALL CAPS line (5+ chars) or ends with ':'
        is_section = (
            (stripped == stripped.upper() and len(stripped) >= 5 and stripped.replace(" ","").replace(":","").replace("-","").isalpha())
            or stripped.endswith(":")
        )
        if is_section:
            add_section(stripped.rstrip(":"))
            i += 1
            continue

        # Bullet: starts with - * •
        if stripped.startswith(("- ", "* ", "• ")):
            add_bullet(stripped[2:])
            i += 1
            continue

        # Regular body
        if stripped:
            add_body(stripped)
        i += 1

    doc.save(out_path)
    print(f"OK:{out_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--title",    required=True)
    parser.add_argument("--subtitle", default="")
    parser.add_argument("--content",  required=True)
    parser.add_argument("--out",      required=True)
    args = parser.parse_args()
    build_doc(args.title, args.subtitle, args.content, args.out)
